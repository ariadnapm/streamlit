import streamlit as st
import fitz
import tempfile
import os
import docx

st.set_page_config(page_title="Drop all Q- and F- numbers", page_icon="📈", layout="wide")
st.sidebar.header("Drop all Q- and F- numbers")

st.markdown(
    """
    <style>
    .reportview-container .main .block-container {
        max-width: 100%;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

container = st.container()

def extract_q_numbers_from_pdf(file_path):
    q_numbers_info = {}
    try:
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            words = text.split()
            for word in words:
                if word.startswith('Q') and word[1:].isdigit():
                    q_number = word
                    if q_number not in q_numbers_info:
                        q_numbers_info[q_number] = {'count': 1, 'pages': [page_num + 1]}
                    else:
                        q_numbers_info[q_number]['count'] += 1
                        q_numbers_info[q_number]['pages'].append(page_num + 1)
    except Exception as e:
        st.error(f"Error extracting Q numbers from PDF: {e}")
    return q_numbers_info

def extract_f_numbers_from_pdf(file_path):
    f_numbers_info = {}
    try:
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            words = text.split()
            for word in words:
                if word.startswith('F-') and word[2:].isdigit():
                    f_number = word
                    if f_number not in f_numbers_info:
                        f_numbers_info[f_number] = {'count': 1, 'pages': [page_num + 1]}
                    else:
                        f_numbers_info[f_number]['count'] += 1
                        f_numbers_info[f_number]['pages'].append(page_num + 1)
    except Exception as e:
        st.error(f"Error extracting F- numbers from PDF: {e}")
    return f_numbers_info

def extract_q_numbers_from_docx(file):
    q_numbers_info = {}
    try:
        doc = docx.Document(file)
        for paragraph in doc.paragraphs:
            for word in paragraph.text.split():
                if word.startswith('Q') and word[1:].isdigit():
                    q_number = word
                    if q_number not in q_numbers_info:
                        q_numbers_info[q_number] = {'count': 1}
                    else:
                        q_numbers_info[q_number]['count'] += 1
    except Exception as e:
        st.error(f"Error extracting Q numbers from DOCX: {e}")
    return q_numbers_info

def extract_f_numbers_from_docx(file):
    f_numbers_info = {}
    try:
        doc = docx.Document(file)
        for paragraph in doc.paragraphs:
            for word in paragraph.text.split():
                if word.startswith('F-') and word[2:].isdigit():
                    f_number = word
                    if f_number not in f_numbers_info:
                        f_numbers_info[f_number] = {'count': 1}
                    else:
                        f_numbers_info[f_number]['count'] += 1
    except Exception as e:
        st.error(f"Error extracting F- numbers from DOCX: {e}")
    return f_numbers_info

def main():
    st.title("Q and F- Number Extractor")
    st.sidebar.title("Q and F- Number Extractor")
    st.sidebar.write("Upload a PDF or DOCX file to extract Q and F- numbers.")

    uploaded_file = st.file_uploader("Upload PDF or DOCX", type=["pdf", "docx"])

    if uploaded_file is not None:
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            temp_file.write(uploaded_file.read())
            file_path = temp_file.name

        if uploaded_file.type == "application/pdf":
            q_numbers_info = extract_q_numbers_from_pdf(file_path)
            f_numbers_info = extract_f_numbers_from_pdf(file_path)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            q_numbers_info = extract_q_numbers_from_docx(file_path)
            f_numbers_info = extract_f_numbers_from_docx(file_path)

        st.write("### Q Numbers extracted:")
        st.table([{'Number': q_number, 'Count': info['count'], 'Pages' if 'pages' in info else 'Pages': ', '.join(map(str, info['pages'])) if 'pages' in info else ''} for q_number, info in q_numbers_info.items()])

        st.write("### F Numbers extracted:")
        st.table([{'Number': f_number, 'Count': info['count'], 'Pages' if 'pages' in info else 'Pages': ', '.join(map(str, info['pages'])) if 'pages' in info else ''} for f_number, info in f_numbers_info.items()])


if __name__ == "__main__":
    main()
